"""
Universal Learning Engine — Teaches the AI agent from ANY source.
Learns from: text, URLs, files, PDFs, code, conversations, corrections,
user feedback, structured data, and its own mistakes.
"""

import os
import re
import json
import hashlib
import fnmatch
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple


class LearningEngine:
    """
    The agent's brain — learns from everything, remembers forever.
    All learned knowledge is stored in both vector memory (semantic search)
    and SQLite (structured queries).
    """

    def __init__(self, vector_memory, database, llm_provider):
        self.vmem = vector_memory
        self.db = database
        self.llm = llm_provider

        # Learning stats (default template)
        self.stats_template = {
            "total_learned": 0,
            "from_text": 0,
            "from_urls": 0,
            "from_files": 0,
            "from_feedback": 0,
            "from_conversations": 0,
            "from_corrections": 0,
            "from_code": 0,
            "from_mistakes": 0,
            "skills_acquired": 0,
        }

    def _load_stats(self, tenant_id: int):
        rows = self.db.search_knowledge(tenant_id, "learning_stats", category="system")
        if rows:
            try:
                # Use tenant specific stats if available
                return json.loads(rows[0]["value"])
            except (json.JSONDecodeError, KeyError):
                pass
        return self.stats_template.copy()

    def _save_stats(self, tenant_id: int, stats: Dict):
        self.db.store_knowledge(tenant_id, "system", "learning_stats",
                                json.dumps(stats), source="learning_engine")

    def _content_hash(self, text: str) -> str:
        return hashlib.md5(text.encode()).hexdigest()[:12]

    # ==================================================
    #  LEARN FROM TEXT (any raw knowledge)
    # ==================================================
    def learn_text(self, tenant_id: int, text: str, topic: str = "general",
                   source: str = "user") -> Dict[str, Any]:
        """Learn from raw text — could be facts, instructions, anything."""
        if not text.strip():
            return {"success": False, "error": "Empty text"}

        doc_id = f"text_{self._content_hash(text)}"

        # Store in vector memory for semantic retrieval
        self.vmem.add(
            tenant_id=tenant_id,
            text=text,
            metadata={"category": "knowledge", "topic": topic, "source": source,
                       "type": "text"},
            doc_id=doc_id
        )

        # Store in structured DB
        self.db.store_knowledge(
            tenant_id=tenant_id,
            category=topic, key=doc_id, value=text,
            confidence=0.9, source=source
        )

        self.db.audit(tenant_id, "learn_text", f"Learned about '{topic}' ({len(text)} chars)",
                       source="learning_engine")

        # Update stats
        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["from_text"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "id": doc_id, "topic": topic,
                "chars": len(text)}

    # ==================================================
    #  LEARN FROM URL (web page content)
    # ==================================================
    def learn_url(self, tenant_id: int, url: str, topic: str = "web") -> Dict[str, Any]:
        """Learn from a web URL — extracts and stores content."""
        try:
            import requests
            resp = requests.get(url, timeout=15, headers={
                "User-Agent": "Mozilla/5.0 (AI Learning Agent)"
            })
            resp.raise_for_status()
            html = resp.text

            # Strip HTML tags (basic extraction)
            text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()

            # Truncate very long pages
            if len(text) > 15000:
                text = text[:15000] + "... [truncated]"

            if len(text) < 50:
                return {"success": False, "error": "Too little content extracted"}

            # Summarize with LLM for better storage
            summary = self.llm.call(
                f"Summarize this web page content in 3-5 key points:\n\n{text[:5000]}",
                system="You are a knowledge extraction assistant. Be concise.",
                max_tokens=500
            )

            doc_id = f"url_{self._content_hash(url)}"

            # Store both raw and summary
            self.vmem.add(
                tenant_id=tenant_id,
                text=f"[Source: {url}]\n{summary}",
                metadata={"category": "knowledge", "topic": topic,
                           "source": url, "type": "url"},
                doc_id=doc_id
            )

            self.vmem.add(
                tenant_id=tenant_id,
                text=text[:5000],
                metadata={"category": "raw_web", "source": url, "type": "url_raw"},
                doc_id=f"{doc_id}_raw"
            )

            self.db.store_knowledge(
                tenant_id,
                category=topic, key=url, value=summary,
                confidence=0.8, source=url
            )

            stats = self._load_stats(tenant_id)
            stats["total_learned"] += 1
            stats["from_urls"] += 1
            self._save_stats(tenant_id, stats)

            return {"success": True, "url": url, "summary": summary[:200],
                    "chars_extracted": len(text)}

        except Exception as e:
            return {"success": False, "error": str(e)}

    # ==================================================
    #  LEARN FROM FILE (txt, py, json, csv, md, etc.)
    # ==================================================
    def learn_file(self, tenant_id: int, filepath: str, topic: str = "file") -> Dict[str, Any]:
        """Learn from a local file — supports text, code, json, csv, markdown."""
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}

        ext = os.path.splitext(filepath)[1].lower()
        filename = os.path.basename(filepath)

        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        except Exception as e:
            return {"success": False, "error": f"Cannot read file: {e}"}

        if not content.strip():
            return {"success": False, "error": "File is empty"}

        # Truncate very large files
        if len(content) > 20000:
            content = content[:20000] + "\n... [truncated]"

        # Extract key info based on file type
        file_type = "text"
        if ext in (".py", ".js", ".ts", ".java", ".cpp", ".c", ".go", ".rs"):
            file_type = "code"
        elif ext == ".json":
            file_type = "data"
            try:
                data = json.loads(content)
                content = f"JSON structure: {json.dumps(data, indent=2)[:5000]}"
            except json.JSONDecodeError:
                pass
        elif ext == ".csv":
            file_type = "data"
            lines = content.split("\n")
            content = f"CSV with {len(lines)} rows. Headers: {lines[0] if lines else 'none'}\nSample:\n" + "\n".join(lines[:10])
        elif ext in (".md", ".txt", ".rst"):
            file_type = "document"

        doc_id = f"file_{self._content_hash(filepath)}"

        self.vmem.add(
            tenant_id=tenant_id,
            text=f"[File: {filename}] ({file_type})\n{content[:8000]}",
            metadata={"category": "knowledge", "topic": topic,
                       "source": filepath, "type": f"file_{file_type}",
                       "filename": filename, "extension": ext},
            doc_id=doc_id
        )

        self.db.store_knowledge(
            tenant_id,
            category=topic, key=filepath,
            value=f"[{file_type}] {filename}: {content[:2000]}",
            confidence=0.85, source=filepath
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["from_files"] += 1
        if file_type == "code":
             stats["from_code"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "file": filename, "type": file_type,
                "chars": len(content)}

    # ==================================================
    #  LEARN FROM PDF (pypdf-powered)
    # ==================================================
    def learn_pdf(self, tenant_id: int, filepath: str, topic: str = "pdf") -> Dict[str, Any]:
        """Ingest a PDF file into vector memory — chunks per page."""
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        filename = os.path.basename(filepath)
        try:
            from pypdf import PdfReader
        except ImportError:
            return {"success": False, "error": "pypdf not installed. Run: pip install pypdf"}

        try:
            reader = PdfReader(filepath)
            total_pages = len(reader.pages)
            chunks_stored = 0
            total_chars = 0

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    continue

                # Chunk large pages into ~2000-char windows
                chunk_size = 2000
                for chunk_idx, start in enumerate(range(0, len(text), chunk_size)):
                    chunk = text[start:start + chunk_size].strip()
                    if len(chunk) < 50:
                        continue
                    doc_id = f"pdf_{self._content_hash(filepath)}_p{page_num}c{chunk_idx}"
                    self.vmem.add(
                        tenant_id=tenant_id,
                        text=f"[PDF: {filename} | Page {page_num}/{total_pages}]\n{chunk}",
                        metadata={"category": "document", "topic": topic,
                                  "source": filepath, "type": "pdf",
                                  "filename": filename, "page": str(page_num)},
                        doc_id=doc_id
                    )
                    chunks_stored += 1
                    total_chars += len(chunk)

            self.db.store_knowledge(
                tenant_id, category=topic, key=filepath,
                value=f"[PDF] {filename}: {total_pages} pages, {chunks_stored} chunks ingested",
                confidence=0.9, source=filepath
            )
            stats = self._load_stats(tenant_id)
            stats["total_learned"] += chunks_stored
            stats["from_files"] += 1
            self._save_stats(tenant_id, stats)

            return {"success": True, "file": filename, "pages": total_pages,
                    "chunks": chunks_stored, "chars": total_chars}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ==================================================
    #  LEARN FROM DOCX (python-docx powered)
    # ==================================================
    def learn_docx(self, tenant_id: int, filepath: str, topic: str = "docx") -> Dict[str, Any]:
        """Ingest a Word (.docx) file into vector memory — chunks per section."""
        if not os.path.exists(filepath):
            return {"success": False, "error": f"File not found: {filepath}"}
        filename = os.path.basename(filepath)
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "python-docx not installed. Run: pip install python-docx"}

        try:
            doc = Document(filepath)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(f"[Table row] {row_text}")

            # Chunk into 2000-char windows
            full_text = "\n".join(paragraphs)
            chunk_size = 2000
            chunks_stored = 0
            total_chars = 0

            for chunk_idx, start in enumerate(range(0, len(full_text), chunk_size)):
                chunk = full_text[start:start + chunk_size].strip()
                if len(chunk) < 50:
                    continue
                doc_id = f"docx_{self._content_hash(filepath)}_c{chunk_idx}"
                self.vmem.add(
                    tenant_id=tenant_id,
                    text=f"[DOCX: {filename} | Chunk {chunk_idx + 1}]\n{chunk}",
                    metadata={"category": "document", "topic": topic,
                              "source": filepath, "type": "docx",
                              "filename": filename, "chunk": str(chunk_idx)},
                    doc_id=doc_id
                )
                chunks_stored += 1
                total_chars += len(chunk)

            self.db.store_knowledge(
                tenant_id, category=topic, key=filepath,
                value=f"[DOCX] {filename}: {len(paragraphs)} paragraphs, {chunks_stored} chunks ingested",
                confidence=0.9, source=filepath
            )
            stats = self._load_stats(tenant_id)
            stats["total_learned"] += chunks_stored
            stats["from_files"] += 1
            self._save_stats(tenant_id, stats)

            return {"success": True, "file": filename, "paragraphs": len(paragraphs),
                    "chunks": chunks_stored, "chars": total_chars}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ==================================================
    #  LEARN FROM DOCUMENT (dispatcher: pdf / docx / txt)
    # ==================================================
    def learn_document(self, tenant_id: int, filepath: str, topic: str = "document") -> Dict[str, Any]:
        """Smart dispatcher — routes .pdf, .docx, or text files to the right learner."""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".pdf":
            return self.learn_pdf(tenant_id, filepath, topic)
        elif ext in (".docx", ".doc"):
            return self.learn_docx(tenant_id, filepath, topic)
        else:
            return self.learn_file(tenant_id, filepath, topic)

    # ==================================================
    #  LEARN FROM DIRECTORY (batch file learning)
    # ==================================================
    def learn_directory(self, tenant_id: int, dirpath: str, extensions: List[str] = None,
                        topic: str = "project") -> Dict[str, Any]:
        """Learn from all files in a directory."""
        if not os.path.isdir(dirpath):
            return {"success": False, "error": f"Directory not found: {dirpath}"}

        if extensions is None:
            extensions = [".py", ".js", ".ts", ".md", ".txt", ".json", ".csv",
                          ".java", ".go", ".rs", ".cpp", ".c", ".yaml", ".yml",
                          ".toml", ".cfg", ".ini", ".html", ".css"]

        learned = []
        errors = []
        # Define ignore patterns for directories and files
        ignore_patterns = [
            "__pycache__", ".git", ".vscode", "venv", "env",
            "agent_db.sqlite", "ultimate_agent_memory.json",
            "hive_mind_central.db", "code_integrity_ledger.json",
            "agent_backups", "agent_modules", "agent_vector_db", "recovery", "screenshots",
        ]
        # File patterns to ignore (e.g., "*.zip", "*.log")
        ignore_file_patterns = ["*.zip", "*.log", "*.tmp"]

        for root, dirs, files in os.walk(dirpath):
            # Filter out ignored directories
            dirs[:] = [d for d in dirs if not d.startswith(".") and d not in ignore_patterns]

            for fname in files:
                ext = os.path.splitext(fname)[1].lower()
                # Check if file extension is in allowed extensions and not in ignore file patterns
                if ext in extensions and not any(fnmatch.fnmatch(fname, pattern) for pattern in ignore_file_patterns):
                    fpath = os.path.join(root, fname)
                    result = self.learn_file(tenant_id, fpath, topic=topic)
                    if result["success"]:
                        learned.append(fname)
                    else:
                        errors.append(f"{fname}: {result['error']}")

        return {"success": True, "files_learned": len(learned),
                "errors": len(errors), "files": learned[:20]}

    # ==================================================
    #  LEARN FROM CONVERSATION (auto-extract knowledge)
    # ==================================================
    def learn_from_conversation(self, tenant_id: int, user_msg: str, agent_msg: str,
                                 session_id: str = "") -> Dict[str, Any]:
        """Auto-extract learnable facts from a conversation turn."""
        extraction = self.llm.call(
            f"Extract any facts, preferences, or instructions from this exchange "
            f"that would be useful to remember. Return as bullet points. "
            f"If nothing worth remembering, say NOTHING.\n\n"
            f"User: {user_msg}\nAssistant: {agent_msg}",
            system="You extract and summarize key knowledge from conversations. Be concise.",
            max_tokens=300
        )

        if "NOTHING" in extraction.upper() or len(extraction.strip()) < 20:
            return {"success": True, "extracted": False}

        doc_id = f"conv_{self._content_hash(user_msg)}"
        self.vmem.add(
            tenant_id=tenant_id,
            text=f"[Learned from conversation]: {extraction}",
            metadata={"category": "conversation_learning", "session": session_id,
                       "type": "conversation"},
            doc_id=doc_id
        )

        self.db.store_knowledge(
            tenant_id,
            category="conversation", key=doc_id, value=extraction,
            confidence=0.7, source=f"session:{session_id}"
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["from_conversations"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "extracted": True, "knowledge": extraction[:200]}

    # ==================================================
    #  LEARN FROM CORRECTION (user corrects the agent)
    # ==================================================
    def learn_correction(self, tenant_id: int, original_response: str, correction: str,
                         context: str = "") -> Dict[str, Any]:
        """Learn from user corrections — prevents repeating mistakes."""
        lesson = (
            f"CORRECTION: When asked about '{context[:100]}', "
            f"I incorrectly said: '{original_response[:200]}'. "
            f"The correct answer is: '{correction}'"
        )

        doc_id = f"fix_{self._content_hash(correction)}"
        self.vmem.add(
            tenant_id=tenant_id,
            text=lesson,
            metadata={"category": "correction", "type": "correction",
                       "priority": "high"},
            doc_id=doc_id
        )

        self.db.store_knowledge(
            tenant_id,
            category="corrections", key=doc_id, value=lesson,
            confidence=1.0, source="user_correction"
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["from_corrections"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "lesson": lesson[:200]}

    # ==================================================
    #  LEARN FROM FEEDBACK (thumbs up/down, ratings)
    # ==================================================
    def learn_feedback(self, tenant_id: int, response: str, rating: int,
                       feedback_text: str = "", context: str = "") -> Dict[str, Any]:
        """Learn from user feedback (1-5 rating)."""
        entry = {
            "rating": rating,
            "response_preview": response[:200],
            "feedback": feedback_text,
            "context": context[:200],
            "timestamp": datetime.now().isoformat(),
        }

        if rating <= 2:
            lesson = f"BAD RESPONSE (rated {rating}/5): '{response[:200]}'"
            if feedback_text:
                lesson += f" Feedback: {feedback_text}"
            self.vmem.add(
                tenant_id=tenant_id,
                text=lesson,
                metadata={"category": "feedback_negative", "type": "feedback",
                           "rating": str(rating)},
            )
            
            stats = self._load_stats(tenant_id)
            stats["from_mistakes"] += 1
            self._save_stats(tenant_id, stats)
        elif rating >= 4:
            self.vmem.add(
                tenant_id=tenant_id,
                text=f"GOOD RESPONSE (rated {rating}/5) for '{context[:100]}': {response[:300]}",
                metadata={"category": "feedback_positive", "type": "feedback",
                           "rating": str(rating)},
            )

        self.db.store_knowledge(
            tenant_id,
            category="feedback", key=f"fb_{self._content_hash(response)}",
            value=json.dumps(entry), confidence=1.0, source="user_feedback"
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["from_feedback"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "rating": rating}

    # ==================================================
    #  LEARN A SKILL (structured how-to)
    # ==================================================
    def learn_skill(self, tenant_id: int, skill_name: str, description: str,
                    steps: List[str], examples: List[str] = None) -> Dict[str, Any]:
        """Learn a new skill with structured steps."""
        skill_doc = f"SKILL: {skill_name}\n"
        skill_doc += f"Description: {description}\n"
        skill_doc += "Steps:\n"
        for i, step in enumerate(steps, 1):
            skill_doc += f"  {i}. {step}\n"
        if examples:
            skill_doc += "Examples:\n"
            for ex in examples:
                skill_doc += f"  - {ex}\n"

        doc_id = f"skill_{self._content_hash(skill_name)}"
        self.vmem.add(
            tenant_id=tenant_id,
            text=skill_doc,
            metadata={"category": "skill", "type": "skill",
                       "skill_name": skill_name},
            doc_id=doc_id
        )

        self.db.store_knowledge(
            tenant_id,
            category="skills", key=skill_name, value=skill_doc,
            confidence=0.95, source="skill_learning"
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        stats["skills_acquired"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "skill": skill_name, "steps": len(steps)}

    # ==================================================
    #  LANGUAGE ACQUISITION (Self-Teaching)
    # ==================================================
    def acquire_language(self, language: str) -> Dict[str, Any]:
        """
        Agent dynamically teaches itself a new programming language.
        Requests the LLM to write a compiler wrapper or REPL script, 
        learns the fundamentals, and saves it as a tool.
        """
        print(f"\n[LANGUAGE ACQUISITION] Initiating self-teaching sequence for '{language}'...")
        
        # 1. Ask LLM to generate a quick study guide and execution wrapper
        prompt = (
            f"I need to teach myself {language}. "
            f"1. Give me a 3 bullet-point summary of its core philosophy.\n"
            f"2. Write a Python script (named run_{language.lower()}.py) that takes a string of {language} code, "
            f"saves it to a temp file, runs it using the standard system compiler/interpreter, and returns the output.\n"
            f"Assume the compiler (e.g., rustc, go run, ruby) is installed on the host system."
        )
        
        response = self.llm.call(prompt, system="You are an expert polyglot programmer.", max_tokens=1500)
        
        # 2. Extract the Python wrapper code (naive extraction)
        python_wrapper = ""
        import re
        code_blocks = re.findall(r"```python(.*?)```", response, re.DOTALL)
        if code_blocks:
            python_wrapper = code_blocks[0].strip()
        else:
             return {"success": False, "error": "Failed to generate execution wrapper", "response": response}
             
        # 3. Save the wrapper to the tools directory
        import os
        os.makedirs("tools", exist_ok=True)
        filename = f"tools/run_{language.lower()}.py"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(python_wrapper)
            
        # 4. Store the knowledge in Vector DB
        self.learn_text(1, f"Learned the basics of {language}. I can now execute it using the {filename} tool.", topic="language_skills", source="self_acquisition")
        
        print(f"  [SUCCESS] Acquired '{language}'. Execution wrapper saved to {filename}.")
        return {"success": True, "language": language, "wrapper": filename}

    # ==================================================
    #  LEARN USER PREFERENCES
    # ==================================================
    def learn_preference(self, tenant_id: int, key: str, value: str) -> Dict[str, Any]:
        """Remember a user preference."""
        self.vmem.add(
            tenant_id=tenant_id,
            text=f"USER PREFERENCE: {key} = {value}",
            metadata={"category": "user_preference", "type": "preference",
                       "pref_key": key},
            doc_id=f"pref_{self._content_hash(key)}"
        )

        self.db.store_knowledge(
            tenant_id,
            category="user_preferences", key=key, value=value,
            confidence=1.0, source="user"
        )

        stats = self._load_stats(tenant_id)
        stats["total_learned"] += 1
        self._save_stats(tenant_id, stats)

        return {"success": True, "preference": key}

    # ==================================================
    #  TEACH (LLM-assisted learning from a topic)
    # ==================================================
    def teach_topic(self, tenant_id: int, topic: str, depth: str = "medium") -> Dict[str, Any]:
        """Ask the LLM to teach itself about a topic and store the knowledge."""
        depth_prompts = {
            "brief": "Give a 3-sentence summary of",
            "medium": "Explain in detail (5-10 key points) about",
            "deep": "Give a comprehensive deep-dive on",
        }

        prompt = f"{depth_prompts.get(depth, depth_prompts['medium'])} {topic}"
        knowledge = self.llm.call(
            prompt,
            system="You are an expert knowledge base. Provide accurate, structured information.",
            max_tokens=1500
        )

        if not knowledge or knowledge.startswith("Error"):
            return {"success": False, "error": knowledge}

        result = self.learn_text(tenant_id, knowledge, topic=topic, source="self_taught")
        result["knowledge_preview"] = knowledge[:300]

        return result

    # ==================================================
    #  RECALL (enhanced retrieval)
    # ==================================================
    def recall(self, tenant_id: int, query: str, n: int = 5, include_corrections: bool = True) -> str:
        """Enhanced recall — includes corrections and high-priority memories."""
        results = self.vmem.search(tenant_id, query, n_results=n)

        # Also search corrections specifically
        if include_corrections:
            corrections = self.vmem.search(tenant_id, query, n_results=2, category="correction")
            results = corrections + results

        if not results:
            return ""

        lines = ["[Recalled knowledge]:"]
        seen = set()
        for r in results:
            text = r["text"][:400]
            if text not in seen:
                seen.add(text)
                lines.append(f"  • {text}")

        return "\n".join(lines[:n + 1])

    # ==================================================
    #  KNOWLEDGE STATS
    # ==================================================
    def get_stats(self, tenant_id: int) -> Dict[str, Any]:
        stats = self._load_stats(tenant_id)
        return {
            **stats,
            "vector_memories": self.vmem.count(), # Global count or should be tenant scoped?
        }

    def what_do_i_know(self, tenant_id: int, topic: str = "") -> List[Dict]:
        """List what the agent knows about a topic."""
        if topic:
            return self.db.search_knowledge(tenant_id, topic)
        rows = self.db.conn.execute(
            "SELECT category, COUNT(*) as count FROM knowledge_base WHERE tenant_id=? GROUP BY category ORDER BY count DESC",
            (tenant_id,)
        ).fetchall()
        return [dict(r) for r in rows]
